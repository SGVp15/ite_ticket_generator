import random

from docx import Document
from docx.shared import Inches

from Word.replace import replace_docx_text
from config import DOCX_TEMPLATE
from qr_code.qr_creator import create_qrcode, get_qrcode_text_from_ticket
from ticket_.ticket import Ticket


def get_random_index_list_quest(list_quest, max_group):
    out = []
    for group in range(0, max_group):
        l = []
        for i, v in enumerate(list_quest):
            v1 = v - 1
            if v1 < len(out) or len(out) == 0:
                l.append(random.randint(0, v1))
            else:
                temp_l = []
                for o in out:
                    temp_l.append(o[i])
                while True:
                    n = random.randint(0, v1)
                    if n not in temp_l:
                        break
                l.append(n)
        out.append(l.copy())
    return out


def create_docx(ticket: Ticket):
    document = Document(DOCX_TEMPLATE)
    # replace_docx_text(document, old_text='Exam', new_text=ticket.exam)
    create_qrcode(text=ticket.ticket_name, filename=ticket.file_qrcode_exam_num)
    create_qrcode(text=get_qrcode_text_from_ticket(ticket), filename=ticket.file_qrcode)

    tables = document.tables
    p = tables[3].rows[0].cells[1].add_paragraph()
    r = p.add_run()
    r.add_picture(ticket.file_qrcode, width=Inches(2))

    p = tables[0].rows[0].cells[1].add_paragraph()
    r = p.add_run()
    r.add_picture(ticket.file_qrcode_exam_num, width=Inches(1))
    # r.add_picture(ticket.file_qrcode, width=Inches(1))

    replace_docx_text(document, old_text='Code', new_text=ticket.ticket_name)
    replace_docx_text(document, old_text='Exam', new_text=ticket.exam_name)
    replace_docx_text(document, old_text='Course', new_text=ticket.full_name_course)
    for i in range(ticket.total_count_of_questions):
        try:
            replace_docx_text(
                document,
                old_text=f'Question_{i + 1}',
                new_text=f'{i + 1}.\t' + ticket.questions[i].text_question
            )
            replace_docx_text(
                document,
                old_text=f'Answer_{i + 1}',
                new_text=ticket.questions[i].answer_text
            )
        except IndexError:
            break
    document.save(ticket.word_path_file)
