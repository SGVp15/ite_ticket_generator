import os
import random
import time
from string import ascii_uppercase, digits

import docx2pdf
from docx import Document
from docx.shared import Inches

import Question
from QR import QRcode, get_qrcode_text_from_ticket
from config import docx_template, full_name_course

from main import create_new_ticket, set_to_json, create_txt
from replace import replace_docx_text


def get_random_index_list_quest(list_quest, max_group):
    out = []
    # print([x - 1 for x in list_quest])
    for group in range(0, max_group):
        l = []
        for i, v in enumerate(list_quest):
            v1 = v - 1
            if v1 < len(out) or len(out) == 0:
                l.append(random.randint(0, v1))
            else:
                temp_l = []
                for o in out:
                    temp_l.append(o[i])
                n = 0
                while True:
                    n = random.randint(0, v1)
                    if n not in temp_l:
                        break
                l.append(n)
        out.append(l.copy())
    return out


def convert_docx_to_pdf(file):
    pdf = f'{os.getcwd()}/{file}.pdf'
    print(pdf)
    docx2pdf.convert(f'./{file}.docx', pdf)
    print(file)
    time.sleep(0.1)


def create_new_tickets(exam: str, num):
    for i in num:
        name = f'{exam}'
        name += f'{i:03d}'
        rand = random.choices(ascii_uppercase, k=3)
        name += ''.join(rand)
        rand = random.choices(digits, k=2)
        name += ''.join(rand)

        questions = create_new_ticket(name)
        # set_to_json(obj=ticket, file_name=f'{name}')
        # create_txt(ticket=ticket, name=f'{name}')
        create_docx(questions=questions, name=f'{name}')


def create_docx(questions: [Question], name):
    document = Document(f'../../{docx_template}')
    # replace_docx_text(document, old_text='Exam', new_text=ticket.exam)

    file_qrcode = f'{name}.png'
    QRcode.create_qrcode(text=get_qrcode_text_from_ticket(questions), filename=file_qrcode)
    file_qrcode_exam_num = f'{name}_exam_num.png'
    QRcode.create_qrcode(text=name, filename=file_qrcode_exam_num)

    tables = document.tables
    p = tables[3].rows[0].cells[1].add_paragraph()
    r = p.add_run()
    r.add_picture(file_qrcode, width=Inches(2))

    p = tables[0].rows[0].cells[1].add_paragraph()
    r = p.add_run()
    r.add_picture(file_qrcode_exam_num, width=Inches(1))

    course = full_name_course[questions[1].exam]

    replace_docx_text(document, old_text='Code', new_text=name)
    replace_docx_text(document, old_text='Exam', new_text=questions[1]['exam'])
    replace_docx_text(document, old_text='Course', new_text=course)
    for i in range(30):
        try:
            replace_docx_text(document, old_text=f'Question_{i + 1}',
                              new_text=f'{i + 1}.\t' + questions[i].text_question)
            replace_docx_text(document, old_text=f'Answer_{i + 1}', new_text=questions[i].Answer)
        except IndexError:
            break
    document.save(f'{name}.docx')
